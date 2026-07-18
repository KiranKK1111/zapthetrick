"""Extract plain text from an uploaded file for chat analysis + RAG.

Handles: pdf, docx, xlsx, json, md/txt/csv, **every text/source-code file**
(decoded as UTF-8), and **archives** (zip / 7z / rar) whose text members are
read out and concatenated. Images are NOT handled here (they go to a vision
model upstream).

Password handling: encrypted pdf/docx/xlsx/zip/7z/rar raise `PasswordRequired`
when no/!wrong password is supplied, so the caller can prompt the user (Claude-
style) and retry with `extract_document_text(..., password=...)`.

`extract_document_text(data, filename, password=None)` returns text, or raises
`UnsupportedDocument` / `PasswordRequired` / `FileTooLarge`.
"""
from __future__ import annotations

import io
import json

from app.resume_parser import _extract_docx, _extract_pdf


class UnsupportedDocument(ValueError):
    """The file type isn't something we can extract text from."""


class FileTooLarge(ValueError):
    """The upload exceeds MAX_UPLOAD_BYTES."""


class PasswordRequired(ValueError):
    """The file is password-protected; supply a password and retry."""

    def __init__(self, filename: str) -> None:
        super().__init__(f"{filename}: password required")
        self.filename = filename


# Hard ceilings. 100 MiB per uploaded file (incl. archives).
MAX_UPLOAD_BYTES = 100 * 1024 * 1024
MAX_EXTRACT_CHARS = 4_000_000
# Per-member cap when reading text out of an archive.
_ARCHIVE_MEMBER_CHARS = 200_000
# Decompression-bomb guards: a ≤100 MB archive can legally expand to many GBs
# (or hold millions of tiny members) → OOM/DoS. Cap the total uncompressed size
# and the member count, checked from the header BEFORE any member is read.
_ARCHIVE_MAX_TOTAL_BYTES = 600 * 1024 * 1024   # 600 MiB uncompressed
_ARCHIVE_MAX_MEMBERS = 10_000

IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".webp", ".gif", ".bmp"}
ARCHIVE_EXTS = {
    ".zip", ".7z", ".rar",                       # multi-file (password-aware)
    ".tar", ".tgz", ".tbz2", ".txz", ".tzst",    # tarballs
    ".gz", ".bz2", ".xz", ".zst", ".lz4", ".br",  # single-stream compressors
}
# Office/document + icon binaries we never raw-decode as text. (docx/xlsx/pdf
# have their own extractors when uploaded directly; inside archives they're
# skipped here. Every other binary — exe/images/archives/etc. — is caught by the
# `_looks_texty` sniff, so it doesn't need listing.)
_BINARY_EXTS = {".docx", ".xlsx", ".pdf", ".pptx", ".ico"}


def _cap(text: str, source: str = "document") -> str:
    if len(text) <= MAX_EXTRACT_CHARS:
        return text
    return (text[:MAX_EXTRACT_CHARS]
            + f"\n\n…[{source} truncated at {MAX_EXTRACT_CHARS:,} characters]…")


def _ext(filename: str) -> str:
    name = (filename or "").lower()
    dot = name.rfind(".")
    return name[dot:] if dot != -1 else ""


def is_image(filename: str) -> bool:
    return _ext(filename) in IMAGE_EXTS


def is_archive(filename: str) -> bool:
    return _ext(filename) in ARCHIVE_EXTS


def _looks_texty(data: bytes) -> bool:
    sample = data[:8192]
    if not sample:
        return True
    if b"\x00" in sample:
        return False
    good = sum(1 for b in sample if b in (9, 10, 13) or 32 <= b <= 126 or b >= 128)
    return good / len(sample) > 0.85


def extract_document_text(data: bytes, filename: str, password: str | None = None) -> str:
    """Return the textual content of a file. Raises UnsupportedDocument /
    PasswordRequired."""
    ext = _ext(filename)
    if ext == ".pdf":
        return _cap(_extract_pdf_pw(data, filename, password), filename)
    if ext == ".docx":
        return _cap(_extract_office(data, filename, password, kind="docx"), filename)
    if ext == ".xlsx":
        return _cap(_extract_office(data, filename, password, kind="xlsx"), filename)
    if ext == ".json":
        return _cap(_extract_json(data), filename)
    if ext == ".pptx":
        return _cap(_extract_pptx(data), filename)
    if ext in (".odt", ".odp", ".ods"):
        return _cap(_extract_odf(data), filename)
    if ext in ARCHIVE_EXTS:
        return _cap(_extract_archive(data, filename, password), filename)
    if ext in (".doc", ".ppt", ".xls"):
        raise UnsupportedDocument(
            f"{filename}: legacy binary Office format isn't supported — save it "
            "as the modern .docx/.pptx/.xlsx."
        )
    # Everything else: decode as text if it looks like text (covers md/txt/csv
    # and every source/config/data file — .py/.java/.cpp/.xml/.properties/…).
    if ext in _BINARY_EXTS:
        raise UnsupportedDocument(f"Unsupported file type: {filename}")
    head = data[: MAX_EXTRACT_CHARS + 1]
    if not _looks_texty(head):
        raise UnsupportedDocument(f"{filename}: not a readable text file.")
    text = head.decode("utf-8", errors="replace")
    return _cap(text, filename) if len(data) > MAX_EXTRACT_CHARS else text.strip()


# --------------------------------------------------------------------------- #
# PDF (encrypted-aware)
# --------------------------------------------------------------------------- #
def _extract_pdf_pw(data: bytes, filename: str, password: str | None) -> str:
    import fitz  # PyMuPDF

    doc = fitz.open(stream=data, filetype="pdf")
    try:
        if doc.needs_pass:
            if not doc.authenticate(password or ""):
                raise PasswordRequired(filename)
        return "\n\n".join(p.get_text() for p in doc)
    finally:
        doc.close()


# --------------------------------------------------------------------------- #
# Office (encrypted-aware via msoffcrypto)
# --------------------------------------------------------------------------- #
def _decrypt_office_if_needed(data: bytes, filename: str, password: str | None) -> bytes:
    """OOXML files (docx/xlsx) are ZIPs; an *encrypted* one is an OLE container.
    Detect that and decrypt with msoffcrypto when a password is given."""
    if data[:4] != b"\xd0\xcf\x11\xe0":  # not an OLE compound file → not encrypted
        return data
    try:
        import msoffcrypto
    except Exception:  # noqa: BLE001 — dep missing
        raise PasswordRequired(filename)
    if not password:
        raise PasswordRequired(filename)
    try:
        of = msoffcrypto.OfficeFile(io.BytesIO(data))
        of.load_key(password=password)
        out = io.BytesIO()
        of.decrypt(out)
        return out.getvalue()
    except Exception:  # noqa: BLE001 — wrong password / bad file
        raise PasswordRequired(filename)


def _extract_office(data: bytes, filename: str, password: str | None, *, kind: str) -> str:
    data = _decrypt_office_if_needed(data, filename, password)
    return _extract_docx_md(data) if kind == "docx" else _extract_xlsx(data)


def _extract_docx_md(data: bytes) -> str:
    """Structure-preserving docx → markdown (2026-07-09): heading styles map
    to #-levels, bold/italic runs to **/*, list paragraphs to bullets or
    numbers, tables to markdown tables — so an uploaded docx that gets edited
    in chat and re-exported keeps its shape instead of flattening to plain
    text. Falls back to the plain extractor on any surprise."""
    import re as _re

    try:
        from docx import Document
        from docx.oxml.ns import qn
        from docx.table import Table
        from docx.text.paragraph import Paragraph

        doc = Document(io.BytesIO(data))

        def _runs_md(p) -> str:
            out: list[str] = []
            for r in p.runs:
                t = r.text
                if not t:
                    continue
                if r.bold and r.italic:
                    t = f"***{t}***"
                elif r.bold:
                    t = f"**{t}**"
                elif r.italic:
                    t = f"*{t}*"
                out.append(t)
            return ("".join(out) or p.text or "").strip()

        def _is_list(p) -> bool:
            try:
                ppr = p._p.pPr  # noqa: SLF001 — python-docx has no public API
                return ppr is not None and ppr.numPr is not None
            except Exception:  # noqa: BLE001
                return False

        def _para_md(p) -> str | None:
            text = _runs_md(p)
            if not text:
                return None
            style = ((p.style.name if p.style is not None else "") or "")
            m = _re.match(r"heading\s+(\d)", style.lower())
            if m:
                return "#" * min(6, int(m.group(1))) + " " + (p.text or "").strip()
            if "title" == style.lower():
                return "# " + (p.text or "").strip()
            if _is_list(p) or "list" in style.lower():
                return ("1. " if "number" in style.lower() else "- ") + text
            return text

        blocks: list[str] = []
        for child in doc.element.body.iterchildren():
            if child.tag == qn("w:p"):
                md = _para_md(Paragraph(child, doc))
                if md:
                    blocks.append(md)
            elif child.tag == qn("w:tbl"):
                t = Table(child, doc)
                rows: list[str] = []
                for i, row in enumerate(t.rows):
                    cells = [" ".join((c.text or "").split())
                             for c in row.cells]
                    rows.append("| " + " | ".join(cells) + " |")
                    if i == 0:
                        rows.append(
                            "|" + "|".join(" --- " for _ in cells) + "|")
                if rows:
                    blocks.append("\n".join(rows))
        out = "\n\n".join(blocks).strip()
        if out:
            return out
    except Exception:  # noqa: BLE001 — structure pass is best-effort
        pass
    return _extract_docx(data)


def _extract_xlsx(data: bytes) -> str:
    from openpyxl import load_workbook

    wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    blocks: list[str] = []
    for ws in wb.worksheets:
        lines = [f"## Sheet: {ws.title}"]
        for row in ws.iter_rows(values_only=True):
            cells = ["" if c is None else str(c) for c in row]
            if any(c.strip() for c in cells):
                lines.append("\t".join(cells))
        if len(lines) > 1:
            blocks.append("\n".join(lines))
    wb.close()
    return "\n\n".join(blocks).strip()


def _extract_json(data: bytes) -> str:
    if len(data) > MAX_EXTRACT_CHARS:
        return data[: MAX_EXTRACT_CHARS + 1].decode("utf-8", errors="replace")
    raw = data.decode("utf-8", errors="replace")
    try:
        return json.dumps(json.loads(raw), indent=2, ensure_ascii=False)
    except Exception:  # noqa: BLE001
        return raw.strip()


def _extract_pptx(data: bytes) -> str:
    """PowerPoint text — slides are XML in the .pptx zip; pull the <a:t> runs.
    Dependency-free (zip + regex), enough for the model to analyze the deck."""
    import html
    import re
    import zipfile

    zf = zipfile.ZipFile(io.BytesIO(data))
    slides = sorted(
        (n for n in zf.namelist()
         if n.startswith("ppt/slides/slide") and n.endswith(".xml")),
        key=lambda n: int("".join(filter(str.isdigit, n.rsplit("/", 1)[-1])) or 0),
    )
    out: list[str] = []
    for i, name in enumerate(slides, 1):
        xml = zf.read(name).decode("utf-8", errors="replace")
        lines = []
        for para in re.split(r"</a:p>", xml):  # one line per paragraph
            runs = re.findall(r"<a:t>(.*?)</a:t>", para, re.DOTALL)
            if runs:
                lines.append(html.unescape("".join(runs)))
        if lines:
            out.append(f"## Slide {i}\n" + "\n".join(lines))
    return "\n\n".join(out).strip()


def _extract_odf(data: bytes) -> str:
    """OpenDocument (.odt/.ods/.odp) text — content.xml in the zip, with tags
    stripped and paragraph/row/cell boundaries kept as whitespace."""
    import html
    import re
    import zipfile

    zf = zipfile.ZipFile(io.BytesIO(data))
    try:
        xml = zf.read("content.xml").decode("utf-8", errors="replace")
    except KeyError:
        return ""
    xml = re.sub(r"</text:p>|</text:h>|</table:table-row>", "\n", xml)
    xml = re.sub(r"<text:tab/>|</table:table-cell>", "\t", xml)
    text = re.sub(r"<[^>]+>", "", xml)
    return html.unescape(text).strip()


# --------------------------------------------------------------------------- #
# Archives — read text members out of zip / 7z / rar
# --------------------------------------------------------------------------- #
def _guard_members(count: int, filename: str) -> None:
    if count > _ARCHIVE_MAX_MEMBERS:
        raise FileTooLarge(
            f"{filename}: archive has {count:,} members "
            f"(limit {_ARCHIVE_MAX_MEMBERS:,}) — refusing to extract."
        )


def _guard_total(total: int, filename: str) -> None:
    if total > _ARCHIVE_MAX_TOTAL_BYTES:
        raise FileTooLarge(
            f"{filename}: archive expands to ~{total / 1024 / 1024:.0f} MB "
            f"uncompressed (limit {_ARCHIVE_MAX_TOTAL_BYTES // 1024 // 1024} MB) "
            "— refusing to extract (possible decompression bomb)."
        )


def _member_text(name: str, content: bytes) -> str | None:
    if _ext(name) in _BINARY_EXTS or not _looks_texty(content):
        return None
    return content[:_ARCHIVE_MEMBER_CHARS].decode("utf-8", errors="replace")


def _assemble_archive(filename: str, names: list[str],
                      members: list[tuple[str, bytes]]) -> str:
    out = [f"# Archive: {filename}", "",
           "Contents:", *[f"  - {n}" for n in names], ""]
    for name, content in members:
        txt = _member_text(name, content)
        if txt is not None and txt.strip():
            out.append(f"\n--- {name} ---\n{txt}")
    return "\n".join(out)


def _extract_archive(data: bytes, filename: str, password: str | None) -> str:
    ext = _ext(filename)
    if ext == ".zip":
        return _extract_zip(data, filename, password)
    if ext == ".7z":
        return _extract_7z(data, filename, password)
    if ext == ".rar":
        return _extract_rar(data, filename, password)
    # tar-family + single-stream compressors (gz/bz2/xz/zst/lz4/br) — no passwords.
    return _extract_tar_or_compressed(data, filename)


def _bounded_decompress(filename: str, opener) -> bytes:
    """Decompress a single-stream file object up to the uncompressed cap.

    `opener` returns a file-like object supporting `.read(n)`. We read one byte
    past the cap; if more is produced it's a bomb → FileTooLarge, without ever
    materializing the full (potentially many-GB) payload."""
    limit = _ARCHIVE_MAX_TOTAL_BYTES
    fobj = opener()
    try:
        out = fobj.read(limit + 1)
    finally:
        try:
            fobj.close()
        except Exception:  # noqa: BLE001
            pass
    if len(out) > limit:
        _guard_total(len(out), filename)
    return out


def _bounded_zstd(filename: str, data: bytes) -> bytes:
    try:
        import pyzstd
        return _bounded_decompress(
            filename, lambda: pyzstd.ZstdFile(io.BytesIO(data), "rb")
        )
    except ImportError:
        import zstandard
        dctx = zstandard.ZstdDecompressor()
        return _bounded_decompress(
            filename, lambda: dctx.stream_reader(io.BytesIO(data))
        )


def _bounded_brotli(filename: str, data: bytes) -> bytes:
    """Brotli has no file wrapper, so decompress incrementally with a cap."""
    import brotli

    d = brotli.Decompressor()
    out = bytearray()
    mv = memoryview(data)
    step = 1 << 20
    for i in range(0, len(mv), step):
        out += d.process(bytes(mv[i:i + step]))
        if len(out) > _ARCHIVE_MAX_TOTAL_BYTES:
            _guard_total(len(out), filename)
    return bytes(out)


def _read_tar(raw: bytes, filename: str, mode: str) -> str:
    import tarfile

    tf = tarfile.open(fileobj=io.BytesIO(raw), mode=mode)
    try:
        file_members = [m for m in tf.getmembers() if m.isfile()]
        _guard_members(len(file_members), filename)
        _guard_total(sum(int(getattr(m, "size", 0) or 0) for m in file_members),
                     filename)
        names = tf.getnames()
        members: list[tuple[str, bytes]] = []
        for m in file_members:
            f = tf.extractfile(m)
            if f is not None:
                members.append((m.name, f.read()))
    finally:
        tf.close()
    return _assemble_archive(filename, names, members)


def _extract_tar_or_compressed(data: bytes, filename: str) -> str:
    low = (filename or "").lower()
    if low.endswith((".tar.zst", ".tzst")):
        return _read_tar(_bounded_zstd(filename, data), filename, "r:")
    if low.endswith((".tar", ".tar.gz", ".tgz", ".tar.bz2", ".tbz2",
                     ".tar.xz", ".txz")):
        return _read_tar(data, filename, "r:*")  # tarfile auto-detects gz/bz2/xz

    # Single-stream compressors → decompress the one inner payload (bounded so
    # a bomb can't blow up memory; see _bounded_decompress).
    import bz2
    import gzip
    import lzma

    raw: bytes | None = None
    inner: str | None = None
    if low.endswith(".gz"):
        raw, inner = _bounded_decompress(
            filename, lambda: gzip.open(io.BytesIO(data), "rb")), filename[:-3]
    elif low.endswith(".bz2"):
        raw, inner = _bounded_decompress(
            filename, lambda: bz2.open(io.BytesIO(data), "rb")), filename[:-4]
    elif low.endswith(".xz"):
        raw, inner = _bounded_decompress(
            filename, lambda: lzma.open(io.BytesIO(data), "rb")), filename[:-3]
    elif low.endswith(".zst"):
        raw, inner = _bounded_zstd(filename, data), filename[:-4]
    elif low.endswith(".lz4"):
        import lz4.frame
        raw, inner = _bounded_decompress(
            filename, lambda: lz4.frame.open(io.BytesIO(data), "rb")), filename[:-4]
    elif low.endswith(".br"):
        raw, inner = _bounded_brotli(filename, data), filename[:-3]
    if raw is None:
        raise UnsupportedDocument(f"Unsupported archive: {filename}")
    inner = inner or "file"
    if inner.lower().endswith(".tar"):
        return _read_tar(raw, filename, "r:")
    txt = _member_text(inner, raw)
    if txt is not None:
        return f"# {filename}  ->  {inner}\n\n{txt}"
    return _assemble_archive(filename, [inner], [(inner, raw)])


def _extract_zip(data: bytes, filename: str, password: str | None) -> str:
    import zipfile

    try:
        zf = zipfile.ZipFile(io.BytesIO(data))
    except Exception:  # noqa: BLE001
        raise UnsupportedDocument(f"{filename}: not a valid zip.")
    infos = [zi for zi in zf.infolist() if not zi.is_dir()]
    _guard_members(len(infos), filename)
    _guard_total(sum(int(getattr(zi, "file_size", 0) or 0) for zi in infos), filename)
    encrypted = any(zi.flag_bits & 0x1 for zi in infos)
    if encrypted and not password:
        raise PasswordRequired(filename)
    pwd = password.encode() if password else None
    members: list[tuple[str, bytes]] = []
    for zi in infos:
        try:
            members.append((zi.filename, zf.read(zi, pwd=pwd)))
        except RuntimeError:  # bad password
            raise PasswordRequired(filename)
        except Exception:  # noqa: BLE001 — skip unreadable member
            pass
    return _assemble_archive(filename, [zi.filename for zi in infos], members)


def _extract_7z(data: bytes, filename: str, password: str | None) -> str:
    try:
        import py7zr
    except Exception:  # noqa: BLE001
        raise UnsupportedDocument(f"{filename}: 7z support unavailable (py7zr not installed).")
    try:
        with py7zr.SevenZipFile(io.BytesIO(data), mode="r", password=password) as z:
            if z.needs_password() and not password:
                raise PasswordRequired(filename)
            names = z.getnames()
            _guard_members(len(names), filename)
            # archiveinfo() reads only the header — check the declared
            # uncompressed total BEFORE readall() pulls everything into memory.
            try:
                _info = z.archiveinfo()
                _guard_total(int(getattr(_info, "uncompressed", 0) or 0), filename)
            except (PasswordRequired, FileTooLarge):
                raise
            except Exception:  # noqa: BLE001 — header total unavailable
                pass
            extracted = z.readall()  # {name: BytesIO}
    except (PasswordRequired, FileTooLarge):
        raise
    except Exception:  # noqa: BLE001 — wrong password or corrupt
        if password is None:
            raise PasswordRequired(filename)
        raise PasswordRequired(filename)
    members = [(n, bio.read()) for n, bio in extracted.items()]
    return _assemble_archive(filename, names, members)


def _configure_rar_tool(rarfile) -> None:
    """Point rarfile at the bundled UnRAR.exe so rar works with zero user setup.

    Looks in the frozen app's data dir (_MEIPASS/tools), next to the exe, and
    the dev source tree (zapthetrick_be/tools). Falls back to any unrar/bsdtar
    on PATH if the bundled tool is somehow absent."""
    import os
    import sys

    dirs: list[str] = []
    if getattr(sys, "frozen", False):
        meipass = getattr(sys, "_MEIPASS", "")
        if meipass:
            dirs.append(os.path.join(meipass, "tools"))
        dirs.append(os.path.join(os.path.dirname(sys.executable), "tools"))
    here = os.path.dirname(os.path.dirname(os.path.dirname(__file__)))
    dirs.append(os.path.join(here, "tools"))
    for d in dirs:
        for name in ("UnRAR.exe", "unrar.exe", "unrar"):
            cand = os.path.join(d, name)
            if os.path.isfile(cand):
                rarfile.UNRAR_TOOL = cand
                return


def _extract_rar(data: bytes, filename: str, password: str | None) -> str:
    try:
        import rarfile
    except Exception:  # noqa: BLE001
        raise UnsupportedDocument(f"{filename}: rar support unavailable.")
    _configure_rar_tool(rarfile)
    try:
        rf = rarfile.RarFile(io.BytesIO(data))
        if rf.needs_password() and not password:
            raise PasswordRequired(filename)
        if password:
            rf.setpassword(password)
        infos = [ri for ri in rf.infolist() if not ri.isdir()]
        _guard_members(len(infos), filename)
        _guard_total(sum(int(getattr(ri, "file_size", 0) or 0) for ri in infos),
                     filename)
        members: list[tuple[str, bytes]] = []
        for ri in infos:
            try:
                members.append((ri.filename, rf.read(ri)))
            except (rarfile.PasswordRequired, rarfile.RarWrongPassword):
                raise PasswordRequired(filename)
            except Exception:  # noqa: BLE001
                pass
        return _assemble_archive(filename, [ri.filename for ri in infos], members)
    except (PasswordRequired, FileTooLarge):
        raise
    except (rarfile.PasswordRequired, rarfile.RarWrongPassword):
        # No/wrong password (incl. header-encrypted rars) -> prompt to retry.
        raise PasswordRequired(filename)
    except rarfile.NeedFirstVolume:
        raise UnsupportedDocument(f"{filename}: multi-part rar — upload the first volume.")
    except Exception as exc:  # noqa: BLE001
        raise UnsupportedDocument(f"{filename}: could not read rar ({exc}).")
