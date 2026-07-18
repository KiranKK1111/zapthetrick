"""
Resume text extraction.

Pulls plain text out of an uploaded PDF or DOCX so the LLM can summarise it
into a structured profile. Anything else (txt, rtf, doc) we treat as raw
UTF-8 text — good enough for the demo, but the front-end restricts the file
chooser to .pdf and .docx anyway.
"""
import io

import pdfplumber
from docx import Document


class UnsupportedResumeFormat(ValueError):
    """Raised when the uploaded file extension is something we cannot parse."""


def extract_text(file_bytes: bytes, filename: str) -> str:
    """Return the textual content of a resume file.

    Raises UnsupportedResumeFormat for anything that is not PDF, DOCX, or
    plain text.
    """
    lower = filename.lower()
    if lower.endswith(".pdf"):
        return _extract_pdf(file_bytes)
    if lower.endswith(".docx"):
        return _extract_docx(file_bytes)
    if lower.endswith(".txt"):
        return file_bytes.decode("utf-8", errors="replace")
    raise UnsupportedResumeFormat(
        f"Unsupported file type: {filename}. Use .pdf, .docx, or .txt."
    )


def _extract_pdf(file_bytes: bytes) -> str:
    parts: list[str] = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            text = page.extract_text() or ""
            if text.strip():
                parts.append(text)
    return "\n".join(parts).strip()


def _extract_docx(file_bytes: bytes) -> str:
    doc = Document(io.BytesIO(file_bytes))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    # Also grab text from tables — common in resumes for skills matrices.
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)
    return "\n".join(parts).strip()
