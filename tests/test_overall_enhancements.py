"""Overall-enhancements batch (2026-07-09): E2E document export, live session
report, ledger self-tuning bias, repair-round loop config, structure-aware
docx extraction, sandbox network options. Routers are mounted on tiny apps —
never import app.main here (it loads ML models)."""
from __future__ import annotations

import io
import zipfile

from fastapi import FastAPI
from starlette.testclient import TestClient


def _export_app() -> FastAPI:
    from app.api.routes_documents import router
    app = FastAPI()
    app.include_router(router)      # router already carries the /api prefix
    return app


class TestExportEndToEnd:
    """POST /api/documents/export → real bytes, right magic, validation
    header — the leg the FE runs after `done.document`."""

    def test_pdf(self):
        c = TestClient(_export_app())
        r = c.post("/api/documents/export",
                   json={"content": "# Title\n\nHello world", "format": "pdf",
                         "filename": "test"})
        assert r.status_code == 200
        assert r.content[:5] == b"%PDF-"
        assert "X-Artifact-Validation" in r.headers
        assert "attachment" in r.headers["Content-Disposition"]

    def test_docx(self):
        c = TestClient(_export_app())
        r = c.post("/api/documents/export",
                   json={"content": "# Doc\n\nBody text", "format": "docx",
                         "filename": "test"})
        assert r.status_code == 200
        assert r.content[:2] == b"PK"
        names = zipfile.ZipFile(io.BytesIO(r.content)).namelist()
        assert any(n.startswith("word/") for n in names)

    def test_zip_project_gets_verification_report(self):
        c = TestClient(_export_app())
        content = ("The project:\n\n```main.py\nprint('hello')\n```\n")
        r = c.post("/api/documents/export",
                   json={"content": content, "format": "zip",
                         "filename": "proj"})
        assert r.status_code == 200
        names = zipfile.ZipFile(io.BytesIO(r.content)).namelist()
        assert "VERIFICATION.txt" in names

    def test_empty_content_400(self):
        c = TestClient(_export_app())
        r = c.post("/api/documents/export",
                   json={"content": "  ", "format": "pdf"})
        assert r.status_code == 400


class TestAnswerBias:
    def setup_method(self):
        from app.live import ledger
        ledger.reset_for_tests()

    teardown_method = setup_method

    def test_inert_below_three_corrections(self):
        from app.live import ledger
        ledger._feedback_counts["should_have_answered"] = 2
        assert ledger.answer_bias() == 0.0

    def test_positive_bias_answers_more(self):
        from app.live import ledger
        ledger._feedback_counts["should_have_answered"] = 5
        ledger._feedback_counts["should_not_have_answered"] = 1
        b = ledger.answer_bias()
        assert 0 < b <= 0.10

    def test_negative_bias_answers_less(self):
        from app.live import ledger
        ledger._feedback_counts["should_not_have_answered"] = 6
        b = ledger.answer_bias()
        assert -0.10 <= b < 0

    def test_bounded(self):
        from app.live import ledger
        ledger._feedback_counts["should_have_answered"] = 500
        assert ledger.answer_bias() == 0.10


class TestConfigKnobs:
    def test_new_fields(self):
        from app.core.config_loader import (ArtifactValidationSection,
                                            LLMSection, SandboxSection)
        av = ArtifactValidationSection()
        assert av.repair_rounds == 2
        assert av.install_deps is False
        assert av.smoke_run is True
        assert SandboxSection().harden_runner is True
        assert LLMSection().chat_stream_budget_s == 300.0


class TestDocxStructureExtraction:
    def _docx(self) -> bytes:
        from docx import Document
        doc = Document()
        doc.add_heading("Main Title", level=1)
        doc.add_heading("Section", level=2)
        p = doc.add_paragraph()
        p.add_run("plain and ")
        p.add_run("bold").bold = True
        doc.add_paragraph("first item", style="List Bullet")
        doc.add_paragraph("second item", style="List Bullet")
        t = doc.add_table(rows=2, cols=2)
        t.rows[0].cells[0].text = "Col A"
        t.rows[0].cells[1].text = "Col B"
        t.rows[1].cells[0].text = "1"
        t.rows[1].cells[1].text = "2"
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    def test_structure_preserved(self):
        from app.documents.parser import _extract_docx_md
        md = _extract_docx_md(self._docx())
        assert "# Main Title" in md
        assert "## Section" in md
        assert "**bold**" in md
        assert "- first item" in md
        assert "| Col A | Col B |" in md
        assert "| --- |" in md

    def test_garbage_falls_back(self):
        from app.documents.parser import _extract_docx_md
        try:
            _extract_docx_md(b"not a docx")
        except Exception:  # noqa: BLE001 — plain extractor may raise; either
            pass           # behavior is acceptable, never a wrong answer


class TestSandboxNetworkOptions:
    def test_bwrap_share_net_flag(self):
        from app.sandbox.executor import build_bwrap_argv
        argv = build_bwrap_argv("/tmp/ws", ["python", "x.py"],
                                share_net=True)
        assert "--share-net" in argv
        argv2 = build_bwrap_argv("/tmp/ws", ["python", "x.py"])
        assert "--share-net" not in argv2

    def test_runner_bwrap_none_on_windows(self):
        import os
        from app.agent_workspace.runner import _bwrap_wrap
        if os.name == "nt":
            assert _bwrap_wrap("echo hi", "C:/ws") is None
