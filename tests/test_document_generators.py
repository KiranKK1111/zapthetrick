"""Generate-and-validate tests for every downloadable document format.

These exercise `app.documents.generators.render_document` end to end for each
format the UI offers — txt, md, csv, json, xlsx, docx, pdf, zip — plus the
source-code-file extractor used for language files (py/js/java/…). Each test
generates real bytes and validates them (magic numbers, re-parsing with the
owning library, structural checks), so a regression in any generator is caught.

Pure module under test (docx/openpyxl/fpdf are imported lazily inside the
generators), so these run offline with no DB or LLM.
"""
from __future__ import annotations

import csv as _csv
import io
import json
import re
import zipfile

import pytest

from app.documents.generators import (
    SUPPORTED_FORMATS,
    UnsupportedFormat,
    media_type,
    normalize_format,
    render_document,
)

# A representative answer: headings, prose, a list, a table, and a code block.
SAMPLE_MD = """\
# Quarterly Report

Some **bold** intro prose explaining the numbers below.

## Highlights

- Revenue up 12%
- Two new regions

| Region | Revenue | Growth |
| --- | --- | --- |
| North | 1200 | 8% |
| South | 900 | 15% |

```python
# app/main.py
def main():
    print("hello")
```
"""

SAMPLE_JSON_MD = (
    "Here is the config you asked for:\n\n"
    "```json\n"
    '{"name": "app", "port": 8000, "debug": true, "tags": ["a", "b"]}\n'
    "```\n"
)


# --------------------------------------------------------------------------
# format registry / normalization
# --------------------------------------------------------------------------
def test_all_expected_formats_supported():
    assert set(SUPPORTED_FORMATS) == {
        "txt", "md", "html", "csv", "json", "xlsx", "docx", "pptx", "pdf",
        "zip", "7z"
    }


@pytest.mark.parametrize(
    "alias,canonical",
    [
        ("markdown", "md"), ("text", "txt"), ("word", "docx"),
        ("excel", "xlsx"), ("spreadsheet", "xlsx"), ("geojson", "json"),
        ("PDF", "pdf"), (".docx", "docx"),
        ("powerpoint", "pptx"), ("ppt", "pptx"), ("slides", "pptx"),
        ("presentation", "pptx"), (".pptx", "pptx"),
    ],
)
def test_aliases_normalize(alias, canonical):
    assert normalize_format(alias) == canonical


def test_unsupported_format_raises():
    with pytest.raises(UnsupportedFormat):
        normalize_format("rtf")


def test_media_types():
    assert media_type("pdf") == "application/pdf"
    assert media_type("json").startswith("application/json")
    assert "spreadsheetml" in media_type("xlsx")
    assert "wordprocessingml" in media_type("docx")


# --------------------------------------------------------------------------
# txt / md
# --------------------------------------------------------------------------
def test_txt_strips_fences_and_markers():
    data, mime, ext = render_document(SAMPLE_MD, "txt")
    text = data.decode("utf-8")
    assert ext == "txt" and mime.startswith("text/plain")
    assert "```" not in text          # fences dropped
    assert "# Quarterly" not in text  # heading marker dropped
    assert "Quarterly Report" in text
    assert "print(" in text           # code text kept


def test_md_roundtrips_content():
    data, _, ext = render_document(SAMPLE_MD, "md")
    assert ext == "md"
    assert data.decode("utf-8").strip().startswith("# Quarterly Report")


# --------------------------------------------------------------------------
# csv
# --------------------------------------------------------------------------
def test_csv_from_markdown_table():
    data, _, ext = render_document(SAMPLE_MD, "csv")
    assert ext == "csv"
    rows = list(_csv.reader(io.StringIO(data.decode("utf-8"))))
    assert rows[0] == ["Region", "Revenue", "Growth"]
    assert ["North", "1200", "8%"] in rows


# --------------------------------------------------------------------------
# json
# --------------------------------------------------------------------------
def test_json_extracts_and_pretty_prints():
    data, mime, ext = render_document(SAMPLE_JSON_MD, "json")
    assert ext == "json" and mime.startswith("application/json")
    obj = json.loads(data)  # must be valid JSON
    assert obj == {"name": "app", "port": 8000, "debug": True,
                   "tags": ["a", "b"]}
    assert b"\n  " in data  # pretty-printed (indented)


def test_json_falls_back_to_lossless_wrapper():
    data, _, _ = render_document("Just prose, no JSON here.", "json")
    obj = json.loads(data)
    assert obj == {"content": "Just prose, no JSON here."}


def test_json_finds_bare_object_without_fence():
    data, _, _ = render_document('result: {"ok": true, "n": 3} done', "json")
    assert json.loads(data) == {"ok": True, "n": 3}


# --------------------------------------------------------------------------
# xlsx — valid OOXML, openable by openpyxl, header + data present
# --------------------------------------------------------------------------
def test_xlsx_is_valid_and_has_data():
    data, mime, ext = render_document(SAMPLE_MD, "xlsx")
    assert ext == "xlsx" and "spreadsheetml" in mime
    assert data[:2] == b"PK"  # zip/OOXML magic
    from openpyxl import load_workbook

    wb = load_workbook(io.BytesIO(data))
    ws = wb.active
    header = [c.value for c in ws[1]]
    assert header == ["Region", "Revenue", "Growth"]
    assert ws[1][0].font.bold  # header row bolded


# --------------------------------------------------------------------------
# docx — valid OOXML, openable by python-docx, text present
# --------------------------------------------------------------------------
def test_docx_is_valid_and_has_content():
    data, mime, ext = render_document(SAMPLE_MD, "docx")
    assert ext == "docx" and "wordprocessingml" in mime
    assert data[:2] == b"PK"
    from docx import Document

    doc = Document(io.BytesIO(data))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "Quarterly Report" in text
    assert len(doc.tables) >= 1
    assert doc.tables[0].rows[0].cells[0].text == "Region"


# --------------------------------------------------------------------------
# pdf — valid header + non-trivial size
# --------------------------------------------------------------------------
def test_pdf_is_valid():
    data, mime, ext = render_document(SAMPLE_MD, "pdf", title="Report")
    assert ext == "pdf" and mime == "application/pdf"
    assert data[:5] == b"%PDF-"
    assert data.rstrip().endswith(b"%%EOF")
    assert len(data) > 1000


def test_pdf_renders_via_pymupdf_if_available():
    data, _, _ = render_document(SAMPLE_MD, "pdf", title="Report")
    fitz = pytest.importorskip("fitz")
    doc = fitz.open(stream=data, filetype="pdf")
    assert doc.page_count >= 1
    assert "Quarterly Report" in doc[0].get_text()
    doc.close()


# --------------------------------------------------------------------------
# zip — matches the shown directory tree, real code where available
# --------------------------------------------------------------------------
ZIP_MD = """\
Here's the project.

```
myapp/
├─ app/
│  └─ main.py
└─ README.md
```

```python
# app/main.py
print("hi")
```
"""


def test_zip_is_valid_and_matches_tree():
    data, mime, ext = render_document(ZIP_MD, "zip", title="myapp")
    assert ext == "zip" and mime == "application/zip"
    assert data[:2] == b"PK"
    zf = zipfile.ZipFile(io.BytesIO(data))
    names = set(zf.namelist())
    assert "app/main.py" in names
    assert "README.md" in names
    # Real code lands in the matching path.
    assert b'print("hi")' in zf.read("app/main.py")


def test_7z_is_valid_and_matches_tree():
    import py7zr

    data, mime, ext = render_document(ZIP_MD, "7z", title="myapp")
    assert ext == "7z" and mime == "application/x-7z-compressed"
    # 7z magic number.
    assert data[:6] == b"7z\xbc\xaf\x27\x1c"
    with py7zr.SevenZipFile(io.BytesIO(data), mode="r") as z:
        names = set(z.getnames())
        contents = z.readall()
    assert "app/main.py" in names
    assert "README.md" in names
    body = contents["app/main.py"].read()
    assert b'print("hi")' in body


def test_7z_alias_normalizes():
    from app.documents.generators import normalize_format
    assert normalize_format("7-zip") == "7z"
    assert normalize_format("sevenz") == "7z"


# --------------------------------------------------------------------------
# Phase 1 — model-driven renderers are a faithful refactor of the legacy ones
# --------------------------------------------------------------------------
_RICH_MD = (
    "# Architecture Overview\n\n"
    "Intro paragraph with **bold** and a point.\n\n"
    "## Design\n\n- first bullet\n- second bullet\n\n"
    "1. step one\n2. step two\n\n> a wise quote\n\n"
    "```python\nprint('hi')\n```\n\n"
    "```mermaid\nA -> B -> C\n```\n\n"
    "| Col A | Col B |\n|---|---|\n| 1 | 2 |\n\n"
    "### Deep Section\n\nBody text here.\n")


# These formats all embed a wall-clock timestamp, at second resolution. The two
# renders below happen microseconds apart, but every so often they straddle a
# second boundary — which used to fail this test spuriously (~1 run in 5) even
# though the CONTENT was byte-identical. Timestamps are not what this test is
# about, so normalize them away rather than asserting on them.
#
#   PDF   — /CreationDate, and /ID, whose hash fpdf DERIVES from that timestamp.
#   DOCX  — a ZIP: per-entry mtimes in the archive headers, plus the
#   PPTX     dcterms:created/modified stamps inside docProps/core.xml.
_PDF_VOLATILE_RE = re.compile(
    rb"/(?:CreationDate|ModDate)\s*\(D:[^)]*\)"
    rb"|/ID\s*\[[^\]]*\]")
_OOXML_VOLATILE_RE = re.compile(
    rb"<dcterms:(created|modified)[^>]*>[^<]*</dcterms:\1>")


def _stable_bytes(fmt: str, data: bytes) -> bytes:
    """Strip embedded wall-clock metadata so the comparison is about content."""
    if fmt == "pdf":
        return _PDF_VOLATILE_RE.sub(b"", data)
    if fmt in ("docx", "pptx"):
        # Compare the archive's CONTENTS, not its bytes: a ZIP stores an mtime
        # per entry, so the container churns every second even when every file
        # inside it is identical.
        with zipfile.ZipFile(io.BytesIO(data)) as z:
            parts = []
            for name in sorted(z.namelist()):
                body = _OOXML_VOLATILE_RE.sub(b"", z.read(name))
                parts.append(name.encode() + b"\0" + body)
            return b"\0\0".join(parts)
    return data


@pytest.mark.parametrize("fmt", ["pdf", "docx", "pptx"])
def test_model_driven_matches_legacy_bytes(fmt, monkeypatch):
    """Rendering from the DocumentModel IR (default ON) must be structurally
    equivalent to the legacy Markdown-tuple renderers — proven here by
    byte-equality on a rich document exercising every block type."""
    from app.core.config_loader import cfg

    # Isolate the variable under test: with Phase-4 enrichment OFF, the model
    # and legacy renderers must be byte-identical. (With auto_structure ON they
    # legitimately diverge — the model path enriches the IR directly while the
    # legacy path round-trips through Markdown — so pin it off here.)
    monkeypatch.setattr(cfg.documents, "auto_structure", False, raising=False)
    monkeypatch.setattr(cfg.documents, "model_driven_render", True,
                        raising=False)
    model_bytes, _, ext = render_document(_RICH_MD, fmt, title="My Doc")
    monkeypatch.setattr(cfg.documents, "model_driven_render", False,
                        raising=False)
    legacy_bytes, _, _ = render_document(_RICH_MD, fmt, title="My Doc")
    assert ext == fmt
    assert len(model_bytes) > 500
    assert _stable_bytes(fmt, model_bytes) == _stable_bytes(fmt, legacy_bytes)


# --------------------------------------------------------------------------
# guards
# --------------------------------------------------------------------------
def test_render_unsupported_raises():
    with pytest.raises(UnsupportedFormat):
        render_document("x", "rtf")
